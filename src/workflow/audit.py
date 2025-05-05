import streamlit as st
import pandas as pd
import logging
import os
import hashlib
import datetime
import fitz  # PyMuPDF
from docx import Document

# Custom imports
from cls.document import PDF
from cls.database import Database
from cls.mailclient import Mailclient
from processing.ocr import create_ocr_reader


# Set up logging
log = logging.getLogger(__name__)


@st.cache_data
def get_emails(excluded_ids: list[int] = None):
    """
    Fetch the emails from the mail client.

    :return: The emails fetched from the mail client.
    """
    if excluded_ids:
        return Mailclient.get_instance().get_mails(excluded_ids)
    else:
        return Mailclient.get_instance().get_mails()


def check_for_documents():
    """
    Function to check if there are any documents already stored on the filesystem.
    """
    pass
    # TODO: Implement this function to check if there are any documents already stored on the filesystem.


def assess_emails(emails: pd.DataFrame):
    """
    Function to assess a set of mails and process them accordingly.

    :param emails: The set of mails to process.
    """
    log.info(f'Processing: {len(emails)} selected documents...')
    mailclient = Mailclient.get_instance()
    ocr_reader = create_ocr_reader(language='de')

    # Iterate over the selected documents
    for email_list_email_id in emails:
        log.debug(f'Processing mail with ID {email_list_email_id}')
        attachments = mailclient.get_attachments(email_list_email_id)

        # Check if attachments are present
        if not attachments:
            log.warning(f'No attachments found for mail with ID {email_list_email_id}')
            # st.error(f'No attachments found for mail with ID {email_id}')
            continue
        elif len(attachments) >= 1:
            log.warning(f'Processing mail with ID {email_list_email_id}')
            # TODO: Remove the warning message in favor of a progress bar or spinner (anything that doesn't bloat the UI)
            # st.warning(f'Processing mail with ID {email_list_email_id}')

            for attachment in attachments:
                attachment_case_id = None
                if attachment.get_attributes('content_type') == 'application/pdf':
                    log.info(f'Processing pdf attachment {attachment.get_attributes("filename")}')

                    # Extract text from the document (this also sets various attributes if detected)
                    attachment.extract_table_data(ocr_reader=ocr_reader)

                    # Check if the document contains a valid BaFin ID by checking if the client_id has been set
                    # during the extraction process (this is done by the extract_table_data() method)
                    if attachment.client_id:
                        log.info(f'Document {attachment.email_id} belongs to client {attachment.client_id}')

                        # Get the stage of the audit case
                        stage = attachment.get_audit_stage()
                        # TODO: Add a check if a audit_case already exists for that client & year combination!?

                        # Check the stage of the audit case
                        match stage:
                            case 1:  # Waiting for documents
                                log.info(f'Document found in mail with email_id: {email_list_email_id} for'
                                         f' client_id: {attachment.client_id}.')
                                # Start the validation process
                                process_audit_case(attachment)

                            case 2:  # Data verification
                                log.info(f'Document with email_id: {email_list_email_id} and'
                                         f' client_id: {attachment.client_id} is already in the database.')

                                # TODO: Add a check if the document already has a mail id attached to it (so we
                                #  don't process the same document every time the mail is fetched)

                                # Start the validation process
                                process_audit_case(attachment)

                            case 3:  # Issuing certificate
                                log.warning(
                                    f"Client with id: {attachment.client_id} has already passed the value check ("
                                    f"is currently in phase 3), but submitted a new document with email_id:"
                                    f" {email_list_email_id}.")

                            case 4:  # Completing process
                                log.warning(f"Client with id: {attachment.client_id} has already been certified"
                                            f" (is currently in phase 4), but submitted a new document with email_id: {email_list_email_id}.")

                            case _:  # Default case
                                log.info(f'No case found for client_id: {attachment.client_id}, adding case for'
                                         f' document with email_id: {email_list_email_id}.')

                                # Initialize the audit case
                                attachment.initialize_audit_case(stage=2)

                                # Start the validation process
                                process_audit_case(attachment)

                    else:
                        log.warning(f'No BaFin ID found in document {attachment.get_attributes("filename")}, '
                                    f'email_id: {attachment.email_id}')
                        # TODO: Add a db table or something to store the documents that didn't have a BaFin ID
                else:
                    log.info(f'Skipping non-pdf attachment {attachment.get_attributes("content_type")}')
                    # TODO: Add a db table or something to store the mail ids / attachments that were skipped


def process_audit_case(document: PDF):
    """
    Function to initialize and validate a new audit case for a document.

    :param document: The document to initialize the audit case for.
    """
    db = Database().get_instance()
    case_id = document.get_audit_case_id()

    if document.compare_values():
        # If the values match, set the stage of the audit case to 3
        db.insert(
            f"""
            UPDATE audit_case
            SET stage = 3
            WHERE client_id = ? AND email_id = ?
            """, (document.client_id, document.email_id))

        # Log to application log and audit log
        log.info(
            f"Client with BaFin ID {document.bafin_id} submitted a VALID document with email_id: {document.email_id}",
            audit_log=True, case_id=case_id)

        # Save the document to the filesystem
        document.store_document(case_id)

        # Get match percentage for audit log
        comparison_df = document.get_value_comparison_table()
        if not comparison_df.empty:
            matches = comparison_df['Match status'].value_counts().get('✅', 0)
            total = len(comparison_df)
            match_percentage = (matches / total) * 100 if total > 0 else 0
            log.info(f"Document verification completed with {match_percentage:.1f}% match ({matches}/{total} fields)",
                     audit_log=True, case_id=case_id)

        # TODO: Do we want to do the following parts fully automatic?
        # Proceed to generate the certificate
        if generate_certificate(case_id, db):
            # Update the audit case stage if the certificate was generated successfully
            db.insert(
                f"""
                UPDATE audit_case
                SET stage = 4
                WHERE client_id = ? AND email_id = ?
                """, (document.client_id, document.email_id))

            # Log to application log and audit log
            log.info(
                f"Certificate generated for client with BaFin ID {document.bafin_id} and email_id: {document.email_id}",
                audit_log=True, case_id=case_id)
        else:
            # Log to application log and audit log
            log.error(f"Failed to generate certificate for client with BaFin ID {document.bafin_id}",
                      audit_log=True, case_id=case_id)
    else:
        # If the values do not match, set the stage of the audit case to 2
        db.insert(
            f"""
            UPDATE audit_case
            SET stage = 2
            WHERE client_id = ?
            """, (document.client_id,))

        # Log to application log and audit log
        log.info(
            f"Client with BaFin ID {document.bafin_id} submitted an INVALID document with email_id: {document.email_id}",
            audit_log=True, case_id=case_id)

        # Save the document to the filesystem
        document.store_document(case_id)

        # Get match percentage for audit log
        comparison_df = document.get_value_comparison_table()
        if not comparison_df.empty:
            matches = comparison_df['Match status'].value_counts().get('✅', 0)
            total = len(comparison_df)
            match_percentage = (matches / total) * 100 if total > 0 else 0
            log.info(f"Document verification failed with only {match_percentage:.1f}% match ({matches}/{total} fields)",
                     audit_log=True, case_id=case_id)
    
    # Process any pending log initializations
    from custom_logger import process_pending_log_initializations
    process_pending_log_initializations(db)


def fetch_new_emails(database: Database = Database.get_instance()) -> pd.DataFrame:
    """
    Function to fetch new emails from the mail client.

    :param database: The database instance to use (optional).
    :return: The new emails fetched from the mail client.
    """
    # TODO: Make sure that a email is not marked as processed unless the process finished successfully 
    #  (e.g. if the app crashes but the mail has already been marked "processed",
    #   then we might run into an issue with emails slipping through without processing!)

    # Check what emails have already been processed
    processed_mails = database.query("SELECT DISTINCT email_id FROM document")

    # Fetch the emails from the mail client
    if not processed_mails:
        log.debug('No mails found in the database, fetching all mails.')
        new_mails = get_emails()
    else:
        new_mails = get_emails(processed_mails)
        log.debug(f'Found a total of {len(processed_mails)} mails already in the database.')

    return new_mails


def generate_certificate(audit_case_id: int, database: Database = Database.get_instance()) -> bool:
    """
    Generate a certificate for a validated audit case.

    This function:
    1. Loads the client information from the database
    2. Fills the certificate template with client and audit case information
    3. Combines the certificate with the first page of the audit document and terms & conditions
    4. Saves the final PDF in the audit case folder

    Args:
        audit_case_id: The ID of the audit case
        database: Optional database instance

    Returns:
        bool: True if certificate was generated successfully, False otherwise
    """
    try:
        # Step 1: Get client and audit case information
        client_info = get_client_info(audit_case_id, database)
        if not client_info:
            log.error(f"Failed to get client information for audit case {audit_case_id}")
            return False

        # Get document information
        document_info = get_document_info(audit_case_id, database)
        if not document_info:
            log.error(f"No document found for audit case {audit_case_id}")
            return False

        document_path = document_info['document_path']

        # Step 2: Create certificate from template
        certificate_docx_path = create_certificate_from_template(client_info, audit_case_id)
        if not certificate_docx_path:
            log.error(f"Failed to create certificate for audit case {audit_case_id}")
            return False

        # Step 3: Convert certificate to PDF
        certificate_pdf_path = convert_docx_to_pdf(certificate_docx_path)
        if not certificate_pdf_path:
            log.error(f"Failed to convert certificate to PDF for audit case {audit_case_id}")
            return False

        # Step 4: Extract first page from audit document
        first_page_path = extract_first_page(document_path, audit_case_id)
        if not first_page_path:
            log.error(f"Failed to extract first page from document for audit case {audit_case_id}")
            return False

        # Step 5: Combine PDFs
        terms_conditions_path = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "terms_conditions.pdf")
        if not os.path.exists(terms_conditions_path):
            log.error(f"Terms and conditions file not found at {terms_conditions_path}")
            return False

        combined_pdf_path = combine_pdfs(certificate_pdf_path, first_page_path, terms_conditions_path, audit_case_id)
        if not combined_pdf_path:
            log.error(f"Failed to combine PDFs for audit case {audit_case_id}")
            return False

        # Step 6: Update the database to record that the certificate was generated
        success = update_audit_case(audit_case_id, combined_pdf_path, database)
        if not success:
            log.error(f"Failed to update audit case record for {audit_case_id}")
            return False

        log.info(f"Successfully generated certificate for audit case {audit_case_id}")
        return True

    except Exception as e:
        log.error(f"Error generating certificate for audit case {audit_case_id}: {str(e)}")
        return False


def get_client_info(audit_case_id: int, database: Database) -> dict | None:
    """
    Get client information for the specified audit case.

    Args:
        audit_case_id: The ID of the audit case
        database: Database instance

    Returns:
        dict: Client information or None if not found
    """
    try:
        # Query to get client information
        result = database.query("""
            SELECT 
                c.id as client_id,
                c.institute,
                c.bafin_id,
                c.address,
                c.city,
                a.created_at,
                a.last_updated_at
            FROM 
                audit_case a
            JOIN 
                client c ON a.client_id = c.id
            WHERE 
                a.id = ?
        """, (audit_case_id,))

        if not result or not result[0]:
            log.warning(f"No client information found for audit case {audit_case_id}")
            return None

        # Create a dictionary with client information
        client_info = {
            'client_id': result[0][0],
            'institute': result[0][1],
            'bafin_id': result[0][2],
            'address': result[0][3],
            'city': result[0][4],
            'created_at': result[0][5],
            'validation_date': result[0][6]
        }

        return client_info

    except Exception as e:
        log.error(f"Error getting client information: {str(e)}")
        return None


def get_document_info(audit_case_id: int, database: Database) -> dict | None:
    """
    Get document information for the specified audit case.

    Args:
        audit_case_id: The ID of the audit case
        database: Database instance

    Returns:
        dict: Document information or None if not found
    """
    try:
        # Query to get document information
        result = database.query("""
            SELECT 
                document_path,
                document_filename
            FROM 
                document
            WHERE 
                audit_case_id = ?
            ORDER BY 
                processing_date DESC
            LIMIT 1
        """, (audit_case_id,))

        if not result or not result[0]:
            log.warning(f"No document found for audit case {audit_case_id}")
            return None

        # Create a dictionary with document information
        document_info = {
            'document_path': result[0][0][:-4] + "pdf",  # TODO: This is a workaround and should be fixed!
            'document_filename': result[0][1]
        }

        return document_info

    except Exception as e:
        log.error(f"Error getting document information: {str(e)}")
        return None


def create_certificate_from_template(client_info: dict, audit_case_id: int) -> str | None:
    """
    Create a certificate by filling in the template with client information.

    Args:
        client_info: Dictionary containing client information
        audit_case_id: The ID of the audit case

    Returns:
        str: Path to the created certificate file or None if failed
    """
    try:
        # Get the template path
        template_path = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "certificate_template.docx")

        # Check if template exists
        if not os.path.exists(template_path):
            log.error(f"Certificate template not found at {template_path}")
            return None

        # Load the template document
        doc = Document(template_path)

        # Get current date and year
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        current_year = datetime.datetime.now().year

        # Get validation date from client info or use current date
        validation_date = client_info.get('validation_date', current_date)
        if isinstance(validation_date, str) and 'T' in validation_date:
            # Parse ISO format date
            validation_date = datetime.datetime.fromisoformat(validation_date).strftime("%B %d, %Y")

        # Define replacements for the template
        replacements = {
            "[DATE]": current_date,
            "[YEAR]": str(current_year),
            "[BAFIN_ID]": str(client_info['bafin_id']),
            "[INSTITUTE_NAME]": client_info['institute'],
            "[INSTITUTE_ADDRESS]": client_info['address'],
            "[INSTITUTE_CITY]": client_info['city'],
            "[FISCAL_YEAR_END]": f"December 31, {current_year-1}",  # Assuming fiscal year is previous calendar year
            "[VALIDATION_DATE]": validation_date
        }

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Save the certificate
        certificate_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(certificate_dir, exist_ok=True)

        certificate_path = os.path.join(certificate_dir, f"certificate_{audit_case_id}.docx")
        doc.save(certificate_path)

        log.info(f"Certificate created at {certificate_path}")
        return certificate_path

    except Exception as e:
        log.error(f"Error creating certificate from template: {str(e)}")
        return None


def convert_docx_to_pdf(docx_path: str) -> str | None:
    """
    Convert a DOCX file to PDF using an appropriate method.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        str: Path to the PDF file or None if failed
    """
    try:
        # Define the PDF output path
        pdf_path = docx_path.replace(".docx", ".pdf")

        # Skip docx2pdf on Linux as it explicitly throws an error
        # Only try docx2pdf on Windows
        if os.name == 'nt':  # Windows systems
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)

                # Check if the conversion was successful
                if os.path.exists(pdf_path):
                    log.info(f"DOCX converted to PDF at {pdf_path} using docx2pdf")
                    return pdf_path
            except ImportError:
                log.warning("docx2pdf not available, trying alternative method")
        else:
            log.info("Skipping docx2pdf on non-Windows system")

        # Try using LibreOffice if available
        try:
            import subprocess
            # Get the directory of the docx file
            output_dir = os.path.dirname(docx_path)
            # Get just the filename without path
            filename = os.path.basename(docx_path)

            # Run LibreOffice conversion
            libreoffice_process = subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                docx_path
            ], check=False, capture_output=True)

            # Check if the conversion was successful
            if os.path.exists(pdf_path) and libreoffice_process.returncode == 0:
                log.info(f"DOCX converted to PDF at {pdf_path} using LibreOffice")
                return pdf_path
            else:
                log.warning(f"LibreOffice conversion failed: {libreoffice_process.stderr.decode()}")
        except (ImportError, FileNotFoundError, subprocess.SubprocessError) as e:
            log.warning(f"LibreOffice conversion not available: {str(e)}, trying alternative method")

        # If all else fails, generate a simple PDF from scratch
        log.info("Falling back to ReportLab for PDF generation")
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        # Parse the DOCX to extract text
        doc = Document(docx_path)
        document_text = []
        for para in doc.paragraphs:
            document_text.append(para.text)

        # Create a PDF with the extracted text
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()

        # Build the PDF content
        pdf_content = []
        for text in document_text:
            if text.strip():
                if text.startswith('#'):
                    # Heading style for lines starting with #
                    pdf_content.append(Paragraph(text.replace('#', '').strip(), styles['Heading1']))
                elif text.startswith('**') and text.endswith('**'):
                    # Bold style for text between **
                    pdf_content.append(Paragraph(f"<b>{text.strip('*')}</b>", styles['Normal']))
                else:
                    pdf_content.append(Paragraph(text, styles['Normal']))
                pdf_content.append(Spacer(1, 12))

        # Build the PDF
        pdf_doc.build(pdf_content)

        log.info(f"DOCX converted to PDF at {pdf_path} using reportlab")
        return pdf_path

    except Exception as e:
        log.error(f"Error converting DOCX to PDF: {str(e)}")
        return None


def extract_first_page(document_path: str, audit_case_id: int) -> str | None:
    """
    Extract the first page from a PDF document.

    Args:
        document_path: Path to the PDF document
        audit_case_id: The ID of the audit case

    Returns:
        str: Path to the extracted first page or None if failed
    """
    try:
        # Open the document
        document = fitz.open(document_path)

        # Check if document has pages
        if document.page_count == 0:
            log.error(f"Document has no pages: {document_path}")
            return None

        # Create a new PDF with just the first page
        new_doc = fitz.open()
        new_doc.insert_pdf(document, from_page=0, to_page=0)

        # Save the new document
        output_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"first_page_{audit_case_id}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        document.close()

        log.info(f"First page extracted to {output_path}")
        return output_path

    except Exception as e:
        log.error(f"Error extracting first page: {str(e)}")
        return None


def combine_pdfs(certificate_path: str, first_page_path: str, terms_path: str, audit_case_id: int) -> str | None:
    """
    Combine the certificate, first page of the audit document, and terms & conditions into a single PDF.

    Args:
        certificate_path: Path to the certificate PDF
        first_page_path: Path to the first page of the audit document
        terms_path: Path to the terms & conditions PDF
        audit_case_id: The ID of the audit case

    Returns:
        str: Path to the combined PDF or None if failed
    """
    try:
        # Check if all input files exist
        for path in [certificate_path, first_page_path, terms_path]:
            if not os.path.exists(path):
                log.error(f"Input file not found: {path}")
                return None

        # Create a new PDF document
        result = fitz.open()

        # Add the certificate
        cert_doc = fitz.open(certificate_path)
        result.insert_pdf(cert_doc)
        cert_doc.close()

        # Add the first page of the audit document
        page_doc = fitz.open(first_page_path)
        result.insert_pdf(page_doc)
        page_doc.close()

        # Add the terms & conditions
        terms_doc = fitz.open(terms_path)
        result.insert_pdf(terms_doc)
        terms_doc.close()

        # Save the combined PDF
        output_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"certificate_complete_{audit_case_id}.pdf")
        result.save(output_path)
        result.close()

        log.info(f"Combined PDF saved to {output_path}")
        return output_path

    except Exception as e:
        log.error(f"Error combining PDFs: {str(e)}")
        return None


def update_audit_case(audit_case_id: int, certificate_path: str, database: Database) -> bool:
    """
    Update the audit case to record that a certificate was generated.

    Args:
        audit_case_id: The ID of the audit case
        certificate_path: Path to the generated certificate
        database: Database instance

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Read the certificate file
        with open(certificate_path, 'rb') as file:
            certificate_content = file.read()

        # Generate a hash for the certificate
        certificate_hash = hashlib.md5(certificate_content).hexdigest()

        # Update the audit case to set stage to 4 (process completion)
        database.query("""
            UPDATE audit_case
            SET 
                stage = 4,
                comments = CASE
                    WHEN comments IS NULL THEN 'Certificate generated'
                    ELSE comments || ' | Certificate generated'
                END
            WHERE id = ?
        """, (audit_case_id,))

        # TODO: Fix the display logic first before uncommenting this
        # Insert record in document table for the certificate
        #database.insert("""
        #    INSERT INTO document (
        #        document_hash,
        #        audit_case_id,
        #        document_filename,
        #        document_path,
        #        processed,
        #        processing_date
        #    ) VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        #""", (
        #    certificate_hash,
        #    audit_case_id,
        #    os.path.basename(certificate_path),
        #    certificate_path,
        #    True
        #))

        log.info(f"Database updated for audit case {audit_case_id}")
        return True

    except Exception as e:
        log.error(f"Error updating database: {str(e)}")
        return False