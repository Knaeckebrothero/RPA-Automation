"""
This module contains functions for processing files.
"""
import os
import io
import numpy as np
import logging
import fitz  # This is PyMuPDF
from PIL import Image
from docx import Document
import datetime


# Set up logging
log  = logging.getLogger(__name__)


def get_images_from_pdf(pdf_bytes: bytes) -> list[np.array]:
    pdf_doc = fitz.open("pdf", pdf_bytes)
    extracted_images = []

    for page_num in range(pdf_doc.page_count):
        try:
            page = pdf_doc[page_num]

            for img in page.get_images(full=True):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                pil_image = Image.open(io.BytesIO(image_bytes))
                extracted_images.append(pil_image)  # stores the image in bytes
        except Exception as e:
            log.error(f"Error extracting image from PDF: {e}")
    return extracted_images


def create_certificate_from_template(client_info: dict, audit_case_id: int) -> str | None:
    """
    Create a certificate by filling in the template with client information.

    :param client_info: Dictionary containing client information
    :param audit_case_id: The ID of the audit case

    :return: Path to the created certificate file or None if failed
    """
    try:
        # Get the template path
        template_path = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "certificate_template.docx")

        # Check if template exists
        if not os.path.exists(template_path):
            log.error(f"Certificate template not found at {template_path}")
            return None

        # Load the template document
        doc = Document(template_path)

        # Get current date and year
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        current_year = datetime.datetime.now().year

        # Get validation date from client info or use current date
        validation_date = client_info.get('validation_date', current_date)
        if isinstance(validation_date, str) and 'T' in validation_date:
            # Parse ISO format date
            validation_date = datetime.datetime.fromisoformat(validation_date).strftime("%B %d, %Y")

        # Define replacements for the template
        replacements = {
            "[DATE]": current_date,
            "[YEAR]": str(current_year),
            "[BAFIN_ID]": str(client_info['bafin_id']),
            "[INSTITUTE_NAME]": client_info['institute'],
            "[INSTITUTE_ADDRESS]": client_info['address'],
            "[INSTITUTE_CITY]": client_info['city'],
            "[FISCAL_YEAR_END]": f"December 31, {current_year-1}",  # Assuming fiscal year is previous calendar year
            "[VALIDATION_DATE]": validation_date
        }

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Save the certificate
        certificate_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(certificate_dir, exist_ok=True)

        certificate_path = os.path.join(certificate_dir, f"certificate_{audit_case_id}.docx")
        doc.save(certificate_path)

        log.info(f"Certificate created at {certificate_path}")
        return certificate_path

    except Exception as e:
        log.error(f"Error creating certificate from template: {str(e)}")
        return None


def convert_docx_to_pdf(docx_path: str) -> str | None:
    """
    Convert a DOCX file to PDF using an appropriate method.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        str: Path to the PDF file or None if failed
    """
    try:
        # Define the PDF output path
        pdf_path = docx_path.replace(".docx", ".pdf")

        # Skip docx2pdf on Linux as it explicitly throws an error
        # Only try docx2pdf on Windows
        if os.name == 'nt':  # Windows systems
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)

                # Check if the conversion was successful
                if os.path.exists(pdf_path):
                    log.info(f"DOCX converted to PDF at {pdf_path} using docx2pdf")
                    return pdf_path
            except ImportError:
                log.warning("docx2pdf not available, trying alternative method")
        else:
            log.info("Skipping docx2pdf on non-Windows system")

        # Try using LibreOffice if available
        try:
            import subprocess
            # Get the directory of the docx file
            output_dir = os.path.dirname(docx_path)
            # Get just the filename without path
            filename = os.path.basename(docx_path)

            # Run LibreOffice conversion
            libreoffice_process = subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                docx_path
            ], check=False, capture_output=True)

            # Check if the conversion was successful
            if os.path.exists(pdf_path) and libreoffice_process.returncode == 0:
                log.info(f"DOCX converted to PDF at {pdf_path} using LibreOffice")
                return pdf_path
            else:
                log.warning(f"LibreOffice conversion failed: {libreoffice_process.stderr.decode()}")
        except (ImportError, FileNotFoundError, subprocess.SubprocessError) as e:
            log.warning(f"LibreOffice conversion not available: {str(e)}, trying alternative method")

        # If all else fails, generate a simple PDF from scratch
        log.info("Falling back to ReportLab for PDF generation")
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        # Parse the DOCX to extract text
        doc = Document(docx_path)
        document_text = []
        for para in doc.paragraphs:
            document_text.append(para.text)

        # Create a PDF with the extracted text
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()

        # Build the PDF content
        pdf_content = []
        for text in document_text:
            if text.strip():
                if text.startswith('#'):
                    # Heading style for lines starting with #
                    pdf_content.append(Paragraph(text.replace('#', '').strip(), styles['Heading1']))
                elif text.startswith('**') and text.endswith('**'):
                    # Bold style for text between **
                    pdf_content.append(Paragraph(f"<b>{text.strip('*')}</b>", styles['Normal']))
                else:
                    pdf_content.append(Paragraph(text, styles['Normal']))
                pdf_content.append(Spacer(1, 12))

        # Build the PDF
        pdf_doc.build(pdf_content)

        log.info(f"DOCX converted to PDF at {pdf_path} using reportlab")
        return pdf_path

    except Exception as e:
        log.error(f"Error converting DOCX to PDF: {str(e)}")
        return None


def extract_first_page(document_path: str, audit_case_id: int) -> str | None:
    """
    Extract the first page from a PDF document.

    Args:
        document_path: Path to the PDF document
        audit_case_id: The ID of the audit case

    Returns:
        str: Path to the extracted first page or None if failed
    """
    try:
        # Open the document
        document = fitz.open(document_path)

        # Check if document has pages
        if document.page_count == 0:
            log.error(f"Document has no pages: {document_path}")
            return None

        # Create a new PDF with just the first page
        new_doc = fitz.open()
        new_doc.insert_pdf(document, from_page=0, to_page=0)

        # Save the new document
        output_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"first_page_{audit_case_id}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        document.close()

        log.info(f"First page extracted to {output_path}")
        return output_path

    except Exception as e:
        log.error(f"Error extracting first page: {str(e)}")
        return None


def combine_pdfs(certificate_path: str, first_page_path: str, terms_path: str, audit_case_id: int) -> str | None:
    """
    Combine the certificate, first page of the audit document, and terms & conditions into a single PDF.

    Args:
        certificate_path: Path to the certificate PDF
        first_page_path: Path to the first page of the audit document
        terms_path: Path to the terms & conditions PDF
        audit_case_id: The ID of the audit case

    Returns:
        str: Path to the combined PDF or None if failed
    """
    try:
        # Check if all input files exist
        for path in [certificate_path, first_page_path, terms_path]:
            if not os.path.exists(path):
                log.error(f"Input file not found: {path}")
                return None

        # Create a new PDF document
        result = fitz.open()

        # Add the certificate
        cert_doc = fitz.open(certificate_path)
        result.insert_pdf(cert_doc)
        cert_doc.close()

        # Add the first page of the audit document
        page_doc = fitz.open(first_page_path)
        result.insert_pdf(page_doc)
        page_doc.close()

        # Add the terms & conditions
        terms_doc = fitz.open(terms_path)
        result.insert_pdf(terms_doc)
        terms_doc.close()

        # Save the combined PDF
        output_dir = os.path.join(os.getenv('FILESYSTEM_PATH', './.filesystem'), "documents", str(audit_case_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"certificate_complete_{audit_case_id}.pdf")
        result.save(output_path)
        result.close()

        log.info(f"Combined PDF saved to {output_path}")
        return output_path

    except Exception as e:
        log.error(f"Error combining PDFs: {str(e)}")
        return None
